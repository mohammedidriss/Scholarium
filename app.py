"""Flask web server for Scholarium - Research Assistant (multi-project)."""

import html as html_module
import io
import json
import os
import re
import shutil
import threading
import uuid
import webbrowser
from datetime import datetime, timedelta

from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename

from evaluator import get_evaluator, clear_evaluator_cache

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100 MB max upload

PROJECTS_DIR = os.path.join(os.path.dirname(__file__), "projects")
PROJECTS_FILE = os.path.join(PROJECTS_DIR, "projects.json")
ALLOWED_EXTENSIONS = {"pdf"}


# ---------------------------------------------------------------------------
# Generic JSON helpers
# ---------------------------------------------------------------------------

def _load_json(path, default=None):
    """Load JSON from path, returning default on any failure."""
    if default is None:
        default = []
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError, OSError):
        return default


def _save_json(path, data):
    """Write data as JSON to path."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        json.dump(data, f, indent=2)


# ---------------------------------------------------------------------------
# Project path helpers
# ---------------------------------------------------------------------------

def project_path(pid, filename):
    """Return absolute path for a file inside a project directory."""
    return os.path.join(PROJECTS_DIR, pid, filename)


def project_docs_dir(pid):
    """Return the documents directory for a project."""
    return os.path.join(PROJECTS_DIR, pid, "documents")


def _load_projects():
    """Load the master projects list."""
    return _load_json(PROJECTS_FILE, default=[])


def _save_projects(projects):
    """Persist the master projects list."""
    _save_json(PROJECTS_FILE, projects)


def _get_project(pid):
    """Find a project by id or return None."""
    for p in _load_projects():
        if p["id"] == pid:
            return p
    return None


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# ---------------------------------------------------------------------------
# Migration: move old single-project data into projects/<default_id>/
# ---------------------------------------------------------------------------

def _migrate_old_data():
    """If an old documents/ folder exists at project root, migrate it into a
    default project under the new multi-project structure."""
    base = os.path.dirname(__file__)
    old_docs = os.path.join(base, "documents")
    if not os.path.isdir(old_docs):
        return

    has_pdfs = any(f.lower().endswith(".pdf") for f in os.listdir(old_docs))
    if not has_pdfs:
        return

    pid = uuid.uuid4().hex[:12]
    now = datetime.now().isoformat()

    proj_dir = os.path.join(PROJECTS_DIR, pid)
    os.makedirs(proj_dir, exist_ok=True)

    # Move documents/
    shutil.move(old_docs, os.path.join(proj_dir, "documents"))

    # Move vectordb/ if it exists
    old_vectordb = os.path.join(base, "vectordb")
    if os.path.isdir(old_vectordb):
        shutil.move(old_vectordb, os.path.join(proj_dir, "vectordb"))

    # Move JSON data files
    json_files = [
        "notes.json", "summaries.json", "reading_status.json",
        "highlights.json", "citations.json", "collections.json",
        "journal.json", "literature_matrix.json", "qa_log.json",
    ]
    for jf in json_files:
        old_path = os.path.join(base, jf)
        if os.path.exists(old_path):
            shutil.move(old_path, os.path.join(proj_dir, jf))

    # Register the new project
    projects = _load_projects()
    projects.append({
        "id": pid,
        "name": "Default Project",
        "description": "Migrated from single-project layout",
        "created_at": now,
        "updated_at": now,
    })
    _save_projects(projects)
    print(f"  Migrated old data into project '{pid}' (Default Project)")


# ===========================================================================
# GLOBAL ROUTES
# ===========================================================================

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/health", methods=["GET"])
def health():
    """Check health of all system services."""
    import sys
    import platform
    import flask as fl
    import requests as req

    services = []

    # 1. Python
    services.append({
        "name": "Python",
        "status": "ok",
        "detail": f"{sys.version.split()[0]} ({platform.machine()})",
    })

    # 2. Flask
    services.append({
        "name": "Flask",
        "status": "ok",
        "detail": f"v{fl.__version__} on port 8080",
    })

    # 3. Ollama server
    try:
        r = req.get("http://localhost:11434/api/tags", timeout=3)
        r.raise_for_status()
        models_data = r.json()
        available_models = [m["name"] for m in models_data.get("models", [])]
        available_base = [m["name"].split(":")[0] for m in models_data.get("models", [])]
        model_count = len(models_data.get("models", []))
        services.append({"name": "Ollama", "status": "ok", "detail": f"Running ({model_count} models)"})
    except Exception:
        available_models = []
        available_base = []
        services.append({"name": "Ollama", "status": "error", "detail": "Not reachable"})

    # 4. Respondent model
    from respondent import MODEL as RESP_MODEL
    if RESP_MODEL in available_models or RESP_MODEL in available_base or f"{RESP_MODEL}:latest" in available_models:
        services.append({"name": "Respondent LLM", "status": "ok", "detail": RESP_MODEL})
    else:
        services.append({"name": "Respondent LLM", "status": "error", "detail": f"{RESP_MODEL} not found"})

    # 5. Judge model
    from judge import _get_available_model
    judge_model = _get_available_model()
    if judge_model:
        services.append({"name": "Judge LLM", "status": "ok", "detail": judge_model})
    else:
        services.append({"name": "Judge LLM", "status": "error", "detail": "No judge model"})

    # 6. Embeddings
    services.append({"name": "Embeddings", "status": "ok", "detail": "all-MiniLM-L6-v2"})

    # 7. Total projects count
    projects = _load_projects()
    services.append({
        "name": "Projects",
        "status": "ok",
        "detail": f"{len(projects)} project(s)",
    })

    return jsonify(services)


@app.route("/api/restart", methods=["POST"])
def restart_server():
    """Restart the Flask server by re-executing the process."""
    import sys
    import signal

    def _restart():
        import time
        import subprocess
        time.sleep(1.5)
        subprocess.Popen(
            [sys.executable] + sys.argv,
            cwd=os.path.dirname(os.path.abspath(__file__)),
        )
        os._exit(0)

    threading.Thread(target=_restart, daemon=True).start()
    return jsonify({"success": True, "message": "Restarting..."})


# ===========================================================================
# PROJECT MANAGEMENT ROUTES
# ===========================================================================

@app.route("/api/projects", methods=["GET"])
def list_projects():
    """List all projects, enriched with live document_count."""
    projects = _load_projects()
    out = []
    for p in projects:
        entry = dict(p)
        docs_dir = project_docs_dir(p["id"])
        doc_count = 0
        if os.path.isdir(docs_dir):
            try:
                doc_count = sum(1 for f in os.listdir(docs_dir) if f.lower().endswith(".pdf"))
            except OSError:
                pass
        entry["document_count"] = doc_count
        out.append(entry)
    return jsonify(out)


@app.route("/api/projects", methods=["POST"])
def create_project():
    """Create a new project with directory structure."""
    data = request.get_json() or {}
    if not data.get("name", "").strip():
        return jsonify({"error": "Project name is required"}), 400

    pid = uuid.uuid4().hex[:12]
    now = datetime.now().isoformat()

    proj = {
        "id": pid,
        "name": data["name"].strip(),
        "description": data.get("description", "").strip(),
        "judge_enabled": bool(data.get("judge_enabled", True)),
        "created_at": now,
        "updated_at": now,
    }

    # Create directory structure
    os.makedirs(os.path.join(PROJECTS_DIR, pid, "documents"), exist_ok=True)

    projects = _load_projects()
    projects.append(proj)
    _save_projects(projects)
    return jsonify(proj), 201


@app.route("/api/projects/<pid>", methods=["PUT"])
def update_project(pid):
    """Update project name and/or description."""
    data = request.get_json() or {}
    projects = _load_projects()
    for p in projects:
        if p["id"] == pid:
            description_changed = False
            if "name" in data:
                p["name"] = data["name"].strip()
                description_changed = True  # name changes also affect the semantic query
            if "description" in data:
                if p.get("description") != data["description"].strip():
                    description_changed = True
                p["description"] = data["description"].strip()
            if "judge_enabled" in data:
                p["judge_enabled"] = bool(data["judge_enabled"])
            p["updated_at"] = datetime.now().isoformat()
            _save_projects(projects)

            # If the description changed, recompute relevance for all docs in background
            if description_changed:
                docs_dir = project_docs_dir(pid)
                if os.path.isdir(docs_dir):
                    fnames = [f for f in os.listdir(docs_dir) if f.lower().endswith(".pdf")]
                    if fnames:
                        _trigger_relevance_score(pid, sorted(fnames))

            return jsonify(p)
    return jsonify({"error": "Project not found"}), 404


@app.route("/api/projects/<pid>", methods=["DELETE"])
def delete_project(pid):
    """Delete a project, its directory, and clear evaluator cache."""
    projects = _load_projects()
    filtered = [p for p in projects if p["id"] != pid]
    if len(filtered) == len(projects):
        return jsonify({"error": "Project not found"}), 404

    # Remove project directory
    proj_dir = os.path.join(PROJECTS_DIR, pid)
    if os.path.isdir(proj_dir):
        shutil.rmtree(proj_dir)

    # Clear evaluator cache for this project
    clear_evaluator_cache(pid)

    _save_projects(filtered)
    return jsonify({"success": True, "message": f"Project {pid} deleted"})


# ===========================================================================
# PROJECT-SCOPED DATA ROUTES
# ===========================================================================

# --- Query ---

@app.route("/api/projects/<pid>/query", methods=["POST"])
def project_query(pid):
    """Process a user question through the full RAG pipeline for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json()
    if not data or not data.get("question", "").strip():
        return jsonify({"error": "No question provided"}), 400

    question = data["question"].strip()
    document_filename = (data.get("document_filename") or "").strip() or None

    # If frontend didn't explicitly pass document_filename, try to detect one
    # from a pattern like: "Regarding the document \"xyz.pdf\": ..."
    if not document_filename:
        m = re.match(r'^\s*Regarding the document ["\']([^"\']+\.pdf)["\']\s*:\s*(.*)$', question, re.DOTALL | re.IGNORECASE)
        if m:
            document_filename = m.group(1)
            question = m.group(2).strip() or "Summarize the key points of this document."

    ev = get_evaluator(pid)
    result = ev.process_query(question, document_filename=document_filename)
    return jsonify(result)


# --- Parse PDF ---

@app.route("/api/projects/<pid>/parse-pdf", methods=["POST"])
def project_parse_pdf(pid):
    """Extract text content from an uploaded PDF."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Only PDF files are allowed"}), 400

    import fitz
    import tempfile
    try:
        pdf_bytes = file.read()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name
        doc = fitz.open(tmp_path)
        page_count = len(doc)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        os.unlink(tmp_path)
        return jsonify({
            "success": True,
            "filename": file.filename,
            "text": text.strip(),
            "page_count": page_count,
        })
    except Exception as e:
        return jsonify({"error": f"Failed to parse PDF: {str(e)}"}), 400


# --- Upload ---

@app.route("/api/projects/<pid>/upload", methods=["POST"])
def project_upload(pid):
    """Upload and index a PDF document into a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Only PDF files are allowed"}), 400

    filename = secure_filename(file.filename)
    docs_dir = project_docs_dir(pid)
    os.makedirs(docs_dir, exist_ok=True)
    filepath = os.path.join(docs_dir, filename)
    file.save(filepath)

    ev = get_evaluator(pid)
    chunks_added = ev.index_single_document(filepath)

    # Auto-trigger summarization + matrix entry in background (non-blocking)
    _auto_process_uploaded_doc(pid, filename)

    return jsonify({
        "success": True,
        "filename": filename,
        "chunks_added": chunks_added,
        "auto_processing": True,
        "message": f"Indexed {filename}: {chunks_added} chunks added"
            if chunks_added > 0
            else f"{filename} was already indexed",
    })


# --- Index ---

@app.route("/api/projects/<pid>/index", methods=["POST"])
def project_reindex(pid):
    """Re-index all documents in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    ev = get_evaluator(pid)
    result = ev.index_documents()
    return jsonify(result)


# --- Status ---

@app.route("/api/projects/<pid>/status", methods=["GET"])
def project_status(pid):
    """Get current status for a project (doc count, chunk count)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    ev = get_evaluator(pid)
    return jsonify(ev.get_status())


# --- Documents ---

@app.route("/api/projects/<pid>/documents", methods=["GET"])
def project_list_documents(pid):
    """List all PDF documents with metadata for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    import fitz
    docs_dir = project_docs_dir(pid)
    if not os.path.isdir(docs_dir):
        return jsonify([])

    # Load relevance + summaries once
    relevance_data = _load_json(project_path(pid, "relevance.json"), default={"scores": {}})
    scores_map = relevance_data.get("scores", {})
    summaries_map = _load_json(project_path(pid, "summaries.json"), default={})
    # Check which summaries / relevance jobs are currently running
    running_jobs = _list_jobs(pid, running_only=True)
    summarizing = {j["target"] for j in running_jobs if j.get("type") == "summarize"}
    relevance_running = any(j.get("type") == "relevance" for j in running_jobs)

    docs = []
    for fname in sorted(os.listdir(docs_dir)):
        if not fname.lower().endswith(".pdf"):
            continue
        fpath = os.path.join(docs_dir, fname)
        size = os.path.getsize(fpath)
        pages = 0
        word_count = 0
        try:
            d = fitz.open(fpath)
            pages = len(d)
            for p in d:
                word_count += len(p.get_text().split())
            d.close()
        except Exception:
            pass
        reading_min = max(1, round(word_count / 250))
        doc_entry = {
            "filename": fname,
            "size_bytes": size,
            "size_display": f"{size / 1024:.0f} KB" if size < 1048576 else f"{size / 1048576:.1f} MB",
            "pages": pages,
            "word_count": word_count,
            "reading_time": f"{reading_min} min read",
        }
        # Attach relevance score status
        rel = scores_map.get(fname)
        if rel:
            doc_entry["relevance_score"] = rel.get("score")
            doc_entry["relevance_cosine"] = rel.get("cosine")
            doc_entry["relevance_status"] = "done"
        elif relevance_running:
            doc_entry["relevance_status"] = "running"
        else:
            doc_entry["relevance_status"] = "pending"
        # Attach summary status
        if fname in summarizing:
            doc_entry["summary_status"] = "running"
        elif fname in summaries_map:
            doc_entry["summary_status"] = "done"
        else:
            doc_entry["summary_status"] = "none"
        docs.append(doc_entry)

    # Auto-trigger relevance scoring for any docs that don't have a score yet
    # (unless a job is already running)
    if not relevance_running:
        unscored = [d["filename"] for d in docs if d.get("relevance_status") == "pending"]
        if unscored:
            _trigger_relevance_score(pid, unscored)
            # Update status for these in the response so UI shows "running"
            for d in docs:
                if d.get("relevance_status") == "pending":
                    d["relevance_status"] = "running"

    return jsonify(docs)


@app.route("/api/projects/<pid>/documents/<filename>/text", methods=["GET"])
def project_document_text(pid, filename):
    """Return extracted text from a PDF, split by page."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    import fitz
    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        return jsonify({"error": "Document not found"}), 404
    try:
        doc = fitz.open(fpath)
        pages = []
        for i, page in enumerate(doc):
            pages.append({"page": i + 1, "text": page.get_text()})
        doc.close()
        return jsonify({"filename": filename, "pages": pages, "total_pages": len(pages)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<pid>/documents/<filename>/download", methods=["GET"])
def project_document_download(pid, filename):
    """Serve a PDF file for download."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        return jsonify({"error": "Document not found"}), 404
    return send_file(fpath, mimetype="application/pdf", as_attachment=True, download_name=filename)


@app.route("/api/projects/<pid>/documents/<filename>", methods=["DELETE"])
def project_delete_document(pid, filename):
    """Delete a document, purge its sidecar data (summary, matrix, etc.), and re-index."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    safe = secure_filename(filename)
    fpath = os.path.join(project_docs_dir(pid), safe)
    if not os.path.exists(fpath):
        return jsonify({"error": "Document not found"}), 404
    os.remove(fpath)

    # Purge this document's data from all sidecar files
    _purge_doc_sidecars(pid, [safe])

    ev = get_evaluator(pid)
    ev.rag.clear_index()
    ev.index_documents()
    return jsonify({"success": True, "message": f"Deleted {filename} and re-indexed"})


def _purge_doc_sidecars(pid: str, filenames: list[str]):
    """Remove entries for the given filenames from all per-document JSON sidecars."""
    if not filenames:
        return
    fn_set = set(filenames)
    # Filename-keyed dicts
    for sidecar in ("summaries.json", "reading_status.json", "highlights.json", "citations.json"):
        path = project_path(pid, sidecar)
        if not os.path.exists(path):
            continue
        try:
            data_obj = _load_json(path, default={})
            if isinstance(data_obj, dict):
                for fn in fn_set:
                    data_obj.pop(fn, None)
                _save_json(path, data_obj)
        except Exception:
            pass
    # Relevance has nested "scores"
    rel_path = project_path(pid, "relevance.json")
    if os.path.exists(rel_path):
        try:
            rel = _load_json(rel_path, default={"scores": {}})
            scores = rel.get("scores", {})
            for fn in fn_set:
                scores.pop(fn, None)
            _save_json(rel_path, rel)
        except Exception:
            pass
    # Literature matrix has an entries list
    m_path = project_path(pid, "literature_matrix.json")
    if os.path.exists(m_path):
        try:
            matrix = _load_json(m_path, default={"entries": []})
            matrix["entries"] = [e for e in matrix.get("entries", []) if e.get("filename") not in fn_set]
            _save_json(m_path, matrix)
        except Exception:
            pass


@app.route("/api/projects/<pid>/documents/bulk-delete", methods=["POST"])
def project_bulk_delete_documents(pid):
    """Delete multiple documents in one pass (single re-index at the end)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    filenames = data.get("filenames", [])
    if not isinstance(filenames, list) or not filenames:
        return jsonify({"error": "filenames list required"}), 400

    docs_dir = project_docs_dir(pid)
    deleted = []
    errors = []
    for raw_name in filenames:
        safe = secure_filename(raw_name or "")
        if not safe:
            errors.append({"filename": raw_name, "error": "Invalid filename"})
            continue
        fpath = os.path.join(docs_dir, safe)
        if not os.path.exists(fpath):
            errors.append({"filename": safe, "error": "Not found"})
            continue
        try:
            os.remove(fpath)
            deleted.append(safe)
        except OSError as e:
            errors.append({"filename": safe, "error": str(e)})

    # Purge per-document data from all sidecar files (summaries, matrix, etc.)
    _purge_doc_sidecars(pid, deleted)

    # Single re-index at the end
    ev = get_evaluator(pid)
    ev.rag.clear_index()
    ev.index_documents()

    return jsonify({
        "success": True,
        "deleted": deleted,
        "errors": errors,
        "message": f"Deleted {len(deleted)} document(s){f', {len(errors)} errors' if errors else ''} and re-indexed.",
    })


# --- Summaries ---

def _do_summarize_document(pid: str, filename: str) -> dict:
    """Synchronous core summarization logic. Returns the summary dict or raises."""
    import fitz
    import requests as req
    from respondent import MODEL, OLLAMA_URL

    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        raise FileNotFoundError(f"Document not found: {filename}")

    # Extract first ~4 pages
    doc = fitz.open(fpath)
    text = ""
    for i, page in enumerate(doc):
        if i >= 4:
            break
        text += page.get_text()
    doc.close()
    text = text[:8000]

    prompt = (
        f"Analyze the following academic paper excerpt and provide:\n"
        f"1. A concise SUMMARY (3-5 sentences covering the paper's purpose, methodology, and conclusions)\n"
        f"2. KEY FINDINGS (3-6 bullet points of the most important results/contributions)\n"
        f"3. METHODOLOGY (1-2 sentences on the research approach)\n\n"
        f"Format your response EXACTLY as:\n"
        f"SUMMARY:\n<summary text>\n\n"
        f"KEY FINDINGS:\n- <finding 1>\n- <finding 2>\n...\n\n"
        f"METHODOLOGY:\n<methodology text>\n\n"
        f"--- DOCUMENT TEXT ---\n{text}"
    )

    response = req.post(
        OLLAMA_URL,
        json={
            "model": MODEL,
            "prompt": prompt,
            "system": "You are an academic paper analyst. Provide structured analysis.",
            "stream": False,
            "options": {"temperature": 0.2, "num_predict": 1024},
        },
        timeout=300,
    )
    response.raise_for_status()
    raw = response.json().get("response", "").strip()

    result = {
        "filename": filename,
        "raw": raw,
        "summary": "",
        "key_findings": [],
        "methodology": "",
        "generated_at": datetime.now().isoformat(),
    }

    current = ""
    for line in raw.split("\n"):
        if line.strip().upper().startswith("SUMMARY"):
            current = "summary"; continue
        elif line.strip().upper().startswith("KEY FINDINGS"):
            current = "findings"; continue
        elif line.strip().upper().startswith("METHODOLOGY"):
            current = "methodology"; continue

        if current == "summary":
            result["summary"] += line + " "
        elif current == "findings" and line.strip().startswith("-"):
            result["key_findings"].append(line.strip()[1:].strip())
        elif current == "methodology":
            result["methodology"] += line + " "

    result["summary"] = result["summary"].strip()
    result["methodology"] = result["methodology"].strip()
    if not result["summary"]:
        result["summary"] = raw[:500]

    summaries_path = project_path(pid, "summaries.json")
    summaries = _load_json(summaries_path, default={})
    summaries[filename] = result
    _save_json(summaries_path, summaries)
    return result


def _summarize_worker(pid: str, filename: str, job_id: str):
    """Background worker for summarization."""
    try:
        _do_summarize_document(pid, filename)
        _finish_job(pid, job_id)
    except Exception as e:
        _finish_job(pid, job_id, error=str(e))


def _do_matrix_entry(pid: str, filename: str) -> dict | None:
    """Generate a single matrix entry for a document and append to literature_matrix.json.
    Returns the entry or None on failure."""
    import fitz
    import requests as req
    from respondent import MODEL, OLLAMA_URL

    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        return None

    # Check if already in matrix
    m_path = project_path(pid, "literature_matrix.json")
    matrix = _load_json(m_path, default={"generated_at": None, "entries": []})
    if any(e.get("filename") == filename for e in matrix["entries"]):
        return None

    try:
        doc = fitz.open(fpath)
        text = ""
        for i, page in enumerate(doc):
            if i >= 4:
                break
            text += page.get_text()
        doc.close()
        text = text[:8000]
    except Exception:
        return None

    prompt = (
        "Extract the following metadata from this academic paper text. "
        "Respond in JSON format with these exact keys: "
        "title (string), year (string), methodology (string, brief description of research method), "
        "findings (string, key findings summary), sample_size (string, e.g. 'N=150' or 'N/A' if not applicable).\n\n"
        f"--- DOCUMENT TEXT ---\n{text}"
    )

    try:
        response = req.post(
            OLLAMA_URL,
            json={
                "model": MODEL,
                "prompt": prompt,
                "system": "You are an academic paper metadata extractor. Return valid JSON only.",
                "stream": False,
                "options": {"temperature": 0.1, "num_predict": 512},
            },
            timeout=180,
        )
        response.raise_for_status()
        raw = response.json().get("response", "").strip()

        json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
        metadata = json.loads(json_match.group()) if json_match else json.loads(raw)

        entry = {
            "filename": filename,
            "title": metadata.get("title", ""),
            "year": str(metadata.get("year", "")),
            "methodology": metadata.get("methodology", ""),
            "findings": metadata.get("findings", ""),
            "sample_size": metadata.get("sample_size", ""),
        }

        # Persist
        matrix = _load_json(m_path, default={"generated_at": None, "entries": []})
        if not any(e.get("filename") == filename for e in matrix["entries"]):
            matrix["entries"].append(entry)
        matrix["generated_at"] = datetime.now().isoformat()
        _save_json(m_path, matrix)
        return entry
    except Exception:
        return None


def _matrix_entry_worker(pid: str, filename: str, job_id: str):
    """Background worker for a single matrix entry."""
    try:
        _do_matrix_entry(pid, filename)
        _finish_job(pid, job_id)
    except Exception as e:
        _finish_job(pid, job_id, error=str(e))


# --- Document Relevance Scoring ---

def _cosine_similarity(v1, v2):
    """Cosine similarity between two vectors (list of floats)."""
    import math
    dot = sum(a * b for a, b in zip(v1, v2))
    n1 = math.sqrt(sum(a * a for a in v1))
    n2 = math.sqrt(sum(b * b for b in v2))
    if n1 == 0 or n2 == 0:
        return 0.0
    return dot / (n1 * n2)


def _get_project_description_embedding(pid: str):
    """Compute (or return cached) embedding of the project's description + keywords.
    Returns (embedding_list, description_text) or (None, '') if no description."""
    project = _get_project(pid)
    if not project:
        return None, ""
    desc = (project.get("description") or "").strip()
    name = (project.get("name") or "").strip()
    # Use both name and description for richer context
    query_text = f"{name}. {desc}" if desc else name
    if not query_text.strip():
        return None, ""

    from rag_pipeline import get_embedder
    embedder = get_embedder()
    emb = embedder.encode([query_text])[0].tolist()
    return emb, query_text


def _compute_doc_relevance(pid: str, filename: str) -> dict | None:
    """Compute relevance score (0-100) for a single document against the project description.
    Uses the average embedding of the doc's first ~8 chunks (title + abstract region)."""
    proj_emb, query_text = _get_project_description_embedding(pid)
    if proj_emb is None:
        return None

    from rag_pipeline import get_embedder
    embedder = get_embedder()

    # Load chunks for this file directly from the ChromaDB collection
    ev = get_evaluator(pid)
    try:
        result = ev.rag.collection.get(
            where={"source": filename},
            include=["documents", "embeddings", "metadatas"],
        )
    except Exception:
        return None

    documents = result.get("documents")
    embeddings = result.get("embeddings")
    metadatas = result.get("metadatas")

    # ChromaDB may return numpy arrays — check length explicitly
    n_docs = len(documents) if documents is not None else 0
    n_embs = len(embeddings) if embeddings is not None else 0
    if n_docs == 0 or n_embs == 0:
        return None

    # Build (embedding, metadata) pairs and sort by chunk_index
    pairs = []
    for i in range(n_embs):
        emb = embeddings[i]
        meta = metadatas[i] if (metadatas is not None and i < len(metadatas)) else {}
        pairs.append((emb, meta or {}))
    pairs.sort(key=lambda p: (p[1] or {}).get("chunk_index", 0))

    # Take first 8 chunks (title + abstract + intro) — most representative
    head_embs = []
    for e, _m in pairs[:8]:
        if e is not None:
            # Convert numpy array to list for consistent math
            head_embs.append(list(e) if hasattr(e, "__iter__") else e)
    if not head_embs:
        return None

    # Average embedding (element-wise mean)
    dim = len(head_embs[0])
    avg = [sum(float(e[i]) for e in head_embs) / len(head_embs) for i in range(dim)]

    similarity = _cosine_similarity(proj_emb, avg)
    # Map cosine (-1..1) → 0..100 (clamp negatives to 0 for intuitive UI)
    score = max(0.0, min(1.0, similarity)) * 100

    # Best matching chunk score (for "why this score?" analysis)
    chunk_scores = [_cosine_similarity(proj_emb, e) for e in head_embs]
    best_idx = chunk_scores.index(max(chunk_scores)) if chunk_scores else 0

    return {
        "filename": filename,
        "score": round(score, 1),
        "cosine": round(similarity, 4),
        "description_used": query_text[:200],
        "computed_at": datetime.now().isoformat(),
    }


def _relevance_worker(pid: str, filenames: list, job_id: str):
    """Background worker: compute relevance for given filenames, save to relevance.json."""
    try:
        rel_path = project_path(pid, "relevance.json")
        data = _load_json(rel_path, default={"scores": {}, "description_at": None})

        project = _get_project(pid)
        desc = (project.get("description") or "").strip() if project else ""
        desc_hash = str(hash(desc)) if desc else ""

        for fname in filenames:
            entry = _compute_doc_relevance(pid, fname)
            if entry:
                data["scores"][fname] = entry

        data["description_hash"] = desc_hash
        data["updated_at"] = datetime.now().isoformat()
        _save_json(rel_path, data)
        _finish_job(pid, job_id)
    except Exception as e:
        _finish_job(pid, job_id, error=str(e))


def _trigger_relevance_score(pid: str, filenames: list):
    """Kick off relevance scoring for a list of filenames in the background."""
    if not filenames:
        return
    job_id = f"relevance:{datetime.now().strftime('%H%M%S')}"
    _add_job(pid, job_id, "relevance", f"{len(filenames)} doc(s)")
    threading.Thread(
        target=_relevance_worker,
        args=(pid, filenames, job_id),
        daemon=True,
    ).start()


@app.route("/api/projects/<pid>/relevance", methods=["GET"])
def project_get_relevance(pid):
    """Return the relevance scores for all documents in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    data = _load_json(project_path(pid, "relevance.json"), default={"scores": {}})
    return jsonify(data)


@app.route("/api/projects/<pid>/relevance/recompute", methods=["POST"])
def project_recompute_relevance(pid):
    """Recompute relevance scores for all documents in the project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    docs_dir = project_docs_dir(pid)
    if not os.path.isdir(docs_dir):
        return jsonify({"scores": {}, "message": "No documents."})

    filenames = sorted([f for f in os.listdir(docs_dir) if f.lower().endswith(".pdf")])
    if not filenames:
        return jsonify({"scores": {}, "message": "No documents."})

    _trigger_relevance_score(pid, filenames)
    return jsonify({
        "message": f"Relevance scoring started for {len(filenames)} document(s).",
        "count": len(filenames),
    }), 202


def _auto_process_uploaded_doc(pid: str, filename: str):
    """Kick off summarize + matrix entry + relevance scoring in background for a newly-uploaded doc."""
    sum_job_id = f"summarize:{filename}"
    if not _job_exists_running(pid, sum_job_id):
        summaries = _load_json(project_path(pid, "summaries.json"), default={})
        if filename not in summaries:
            _add_job(pid, sum_job_id, "summarize", filename)
            threading.Thread(
                target=_summarize_worker,
                args=(pid, filename, sum_job_id),
                daemon=True,
            ).start()

    matrix_job_id = f"matrix_entry:{filename}"
    if not _job_exists_running(pid, matrix_job_id):
        matrix = _load_json(project_path(pid, "literature_matrix.json"), default={"generated_at": None, "entries": []})
        if not any(e.get("filename") == filename for e in matrix["entries"]):
            _add_job(pid, matrix_job_id, "matrix_entry", filename)
            threading.Thread(
                target=_matrix_entry_worker,
                args=(pid, filename, matrix_job_id),
                daemon=True,
            ).start()

    # Relevance scoring for this single file
    _trigger_relevance_score(pid, [filename])


@app.route("/api/projects/<pid>/documents/<filename>/summarize", methods=["POST"])
def project_summarize_document(pid, filename):
    """Trigger background summarization. Returns immediately (202).
    Poll GET /summaries or GET /jobs to check completion."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        return jsonify({"error": "Document not found"}), 404

    force = request.args.get("force", "").lower() == "true"
    summaries_path = project_path(pid, "summaries.json")
    summaries = _load_json(summaries_path, default={})

    # Cached and not forced → return immediately
    if filename in summaries and not force:
        return jsonify({"cached": True, **summaries[filename]})

    # If a job is already running for this file, just report it
    job_id = f"summarize:{filename}"
    if _job_exists_running(pid, job_id):
        return jsonify({"status": "already_running", "job_id": job_id}), 202

    # Start background job
    _add_job(pid, job_id, "summarize", filename)
    threading.Thread(
        target=_summarize_worker,
        args=(pid, filename, job_id),
        daemon=True,
    ).start()
    return jsonify({"status": "started", "job_id": job_id, "target": filename}), 202


@app.route("/api/projects/<pid>/summaries", methods=["GET"])
def project_get_summaries(pid):
    """Return all cached document summaries for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    summaries_path = project_path(pid, "summaries.json")
    return jsonify(_load_json(summaries_path, default={}))


# --- Notes ---

@app.route("/api/projects/<pid>/notes", methods=["GET"])
def project_get_notes(pid):
    """Return all notes for a project, most recent first."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    notes_path = project_path(pid, "notes.json")
    notes = _load_json(notes_path, default=[])
    notes.sort(key=lambda n: n.get("updated_at", ""), reverse=True)
    return jsonify(notes)


@app.route("/api/projects/<pid>/notes", methods=["POST"])
def project_create_note(pid):
    """Create a new note in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    now = datetime.now().isoformat()
    note = {
        "id": uuid.uuid4().hex[:12],
        "title": data.get("title", "Untitled Note"),
        "content": data.get("content", ""),
        "tags": data.get("tags", []),
        "created_at": now,
        "updated_at": now,
    }
    notes_path = project_path(pid, "notes.json")
    notes = _load_json(notes_path, default=[])
    notes.append(note)
    _save_json(notes_path, notes)
    return jsonify(note), 201


@app.route("/api/projects/<pid>/notes/<nid>", methods=["PUT"])
def project_update_note(pid, nid):
    """Update an existing note."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    notes_path = project_path(pid, "notes.json")
    notes = _load_json(notes_path, default=[])
    for note in notes:
        if note["id"] == nid:
            if "title" in data:
                note["title"] = data["title"]
            if "content" in data:
                note["content"] = data["content"]
            if "tags" in data:
                note["tags"] = data["tags"]
            note["updated_at"] = datetime.now().isoformat()
            _save_json(notes_path, notes)
            return jsonify(note)
    return jsonify({"error": "Note not found"}), 404


@app.route("/api/projects/<pid>/notes/<nid>", methods=["DELETE"])
def project_delete_note(pid, nid):
    """Delete a note."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    notes_path = project_path(pid, "notes.json")
    notes = _load_json(notes_path, default=[])
    filtered = [n for n in notes if n["id"] != nid]
    if len(filtered) == len(notes):
        return jsonify({"error": "Note not found"}), 404
    _save_json(notes_path, filtered)
    return jsonify({"success": True})


@app.route("/api/projects/<pid>/notes/<nid>/export", methods=["GET"])
def project_export_note(pid, nid):
    """Export a note as PDF or DOCX."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    fmt = request.args.get("format", "pdf").lower()
    notes_path = project_path(pid, "notes.json")
    notes = _load_json(notes_path, default=[])
    note = next((n for n in notes if n["id"] == nid), None)
    if note is None:
        return jsonify({"error": "Note not found"}), 404

    if fmt == "pdf":
        return _export_pdf(note)
    elif fmt == "docx":
        return _export_docx(note)
    else:
        return jsonify({"error": "Unsupported format. Use 'pdf' or 'docx'."}), 400


# --- History ---

@app.route("/api/projects/<pid>/history", methods=["GET"])
def project_get_history(pid):
    """Return all past Q&A entries for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    qa_path = project_path(pid, "qa_log.json")
    history = _load_json(qa_path, default=[])
    history.reverse()
    return jsonify(history)


@app.route("/api/projects/<pid>/history", methods=["DELETE"])
def project_clear_history(pid):
    """Clear all conversation history for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    qa_path = project_path(pid, "qa_log.json")
    _save_json(qa_path, [])
    return jsonify({"success": True})


# --- Reading Status ---

@app.route("/api/projects/<pid>/reading-status", methods=["GET"])
def project_get_reading_status(pid):
    """Return reading status for all documents in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    rs_path = project_path(pid, "reading_status.json")
    return jsonify(_load_json(rs_path, default={}))


@app.route("/api/projects/<pid>/reading-status/<filename>", methods=["PUT"])
def project_update_reading_status(pid, filename):
    """Update reading status and/or progress for a document."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    rs_path = project_path(pid, "reading_status.json")
    statuses = _load_json(rs_path, default={})
    entry = statuses.get(filename, {"status": "unread", "progress_pct": 0, "last_opened": None})

    if "status" in data:
        if data["status"] not in ("unread", "reading", "reviewed"):
            return jsonify({"error": "Invalid status. Must be unread, reading, or reviewed."}), 400
        entry["status"] = data["status"]
    if "progress_pct" in data:
        pct = data["progress_pct"]
        if not isinstance(pct, (int, float)) or pct < 0 or pct > 100:
            return jsonify({"error": "progress_pct must be a number between 0 and 100."}), 400
        entry["progress_pct"] = pct
    entry["last_opened"] = datetime.now().isoformat()

    statuses[filename] = entry
    _save_json(rs_path, statuses)
    return jsonify(entry)


# --- Highlights ---

@app.route("/api/projects/<pid>/highlights/<filename>", methods=["GET"])
def project_get_highlights(pid, filename):
    """Get all highlights for a document in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    hl_path = project_path(pid, "highlights.json")
    highlights = _load_json(hl_path, default={})
    return jsonify(highlights.get(filename, []))


@app.route("/api/projects/<pid>/highlights/<filename>", methods=["POST"])
def project_add_highlight(pid, filename):
    """Add a highlight to a document."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    if not data.get("text"):
        return jsonify({"error": "Highlight text is required."}), 400
    if "page" not in data:
        return jsonify({"error": "Page number is required."}), 400

    highlight = {
        "id": uuid.uuid4().hex[:12],
        "text": data["text"],
        "page": data["page"],
        "created_at": datetime.now().isoformat(),
    }

    hl_path = project_path(pid, "highlights.json")
    highlights = _load_json(hl_path, default={})
    if filename not in highlights:
        highlights[filename] = []
    highlights[filename].append(highlight)
    _save_json(hl_path, highlights)
    return jsonify(highlight), 201


@app.route("/api/projects/<pid>/highlights/<filename>/<hid>", methods=["DELETE"])
def project_delete_highlight(pid, filename, hid):
    """Remove a highlight from a document."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    hl_path = project_path(pid, "highlights.json")
    highlights = _load_json(hl_path, default={})
    if filename not in highlights:
        return jsonify({"error": "No highlights found for this document."}), 404
    original_len = len(highlights[filename])
    highlights[filename] = [h for h in highlights[filename] if h["id"] != hid]
    if len(highlights[filename]) == original_len:
        return jsonify({"error": "Highlight not found."}), 404
    _save_json(hl_path, highlights)
    return jsonify({"success": True})


# --- Citations ---

@app.route("/api/projects/<pid>/citations", methods=["GET"])
def project_list_citations(pid):
    """List all citations for a project as a flat array."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    cit_path = project_path(pid, "citations.json")
    citations_dict = _load_json(cit_path, default={})
    # Convert dict to list with "key" field
    out = []
    for key, entry in citations_dict.items():
        item = dict(entry)
        item["key"] = key
        out.append(item)
    # Sort by first author / year
    out.sort(key=lambda c: (
        (c.get("authors") or [""])[0].lower(),
        c.get("year", ""),
    ))
    return jsonify(out)


@app.route("/api/projects/<pid>/citations/<filename>/generate", methods=["POST"])
def project_generate_citation(pid, filename):
    """Extract citation metadata from first page via LLM."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    import fitz
    import requests as req

    fpath = os.path.join(project_docs_dir(pid), secure_filename(filename))
    if not os.path.exists(fpath):
        return jsonify({"error": "Document not found"}), 404

    force = request.args.get("force", "").lower() == "true"
    cit_path = project_path(pid, "citations.json")
    citations = _load_json(cit_path, default={})
    if filename in citations and not force:
        return jsonify(citations[filename])

    try:
        doc = fitz.open(fpath)
        text = doc[0].get_text() if len(doc) > 0 else ""
        doc.close()
    except Exception as e:
        return jsonify({"error": f"Failed to read PDF: {e}"}), 500

    text = text[:4000]

    prompt = (
        "Extract the following metadata from this academic paper text. "
        "Respond in JSON format with these exact keys: "
        "title, authors (as a list of strings), year (as a string), source_info (journal or conference name).\n\n"
        "Extract: title, authors (list), year, journal/conference from this text:\n\n"
        f"{text}"
    )

    try:
        from respondent import MODEL, OLLAMA_URL
        response = req.post(
            OLLAMA_URL,
            json={
                "model": MODEL,
                "prompt": prompt,
                "system": "You are a metadata extractor. Return valid JSON only.",
                "stream": False,
                "options": {"temperature": 0.1, "num_predict": 512},
            },
            timeout=120,
        )
        response.raise_for_status()
        raw = response.json().get("response", "").strip()

        import re
        json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
        if json_match:
            metadata = json.loads(json_match.group())
        else:
            metadata = json.loads(raw)

        entry = {
            "title": metadata.get("title", ""),
            "authors": metadata.get("authors", []),
            "year": str(metadata.get("year", "")),
            "source_info": metadata.get("source_info", ""),
            "generated_at": datetime.now().isoformat(),
        }

        citations[filename] = entry
        _save_json(cit_path, citations)
        return jsonify(entry)

    except Exception as e:
        return jsonify({"error": f"Citation extraction failed: {str(e)}"}), 500


@app.route("/api/projects/<pid>/citations/<filename>/format", methods=["GET"])
def _format_citation(entry, style):
    """Format a citation entry in the given style.

    Returns the formatted string, or None if the style is unsupported.
    """
    style = (style or "apa").lower()
    title = entry.get("title", "Untitled")
    authors = entry.get("authors", [])
    year = entry.get("year", "n.d.")
    source = entry.get("source_info", "")

    authors_str = ", ".join(authors) if authors else "Unknown"

    # Build a Vancouver-style author list: "Smith J, Jones K"
    def _vancouver_authors(author_list):
        out = []
        for a in author_list:
            parts = a.strip().split()
            if len(parts) >= 2:
                last = parts[-1]
                initials = "".join(p[0] for p in parts[:-1] if p)
                out.append(f"{last} {initials}")
            else:
                out.append(a)
        return ", ".join(out) if out else "Unknown"

    # MLA author format: "Smith, John"
    def _mla_authors(author_list):
        if not author_list:
            return "Unknown"
        a = author_list[0].strip().split()
        if len(a) >= 2:
            first_author = f"{a[-1]}, {' '.join(a[:-1])}"
        else:
            first_author = author_list[0]
        if len(author_list) == 1:
            return first_author
        elif len(author_list) == 2:
            return f"{first_author}, and {author_list[1]}"
        else:
            return f"{first_author}, et al."

    if style == "apa":
        return f"{authors_str} ({year}). {title}. {source}."
    if style == "ieee":
        return f'{authors_str}, "{title}," {source}, {year}.'
    if style == "harvard":
        return f"{authors_str} ({year}) '{title}', {source}."
    if style == "mla":
        return f'{_mla_authors(authors)}. "{title}." {source}, {year}.'
    if style == "chicago":
        return f'{authors_str}. "{title}." {source} ({year}).'
    if style == "vancouver":
        return f"{_vancouver_authors(authors)}. {title}. {source}. {year}."
    return None


def project_format_citation(pid, filename):
    """Return a formatted citation string in the requested style."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    style = request.args.get("style", "apa").lower()
    cit_path = project_path(pid, "citations.json")
    citations = _load_json(cit_path, default={})
    if filename not in citations:
        return jsonify({"error": "No citation data found. Generate citation first."}), 404

    formatted = _format_citation(citations[filename], style)
    if formatted is None:
        return jsonify({"error": "Unsupported style. Use apa, ieee, harvard, mla, chicago, or vancouver."}), 400

    return jsonify({"style": style, "citation": formatted})


# --- BibTeX Import/Export ---

def _citation_to_bibtex(key, entry):
    """Convert a citation entry to a BibTeX string."""
    title = entry.get("title", "Untitled").replace("{", "").replace("}", "")
    authors = entry.get("authors", [])
    year = entry.get("year", "")
    source = entry.get("source_info", "")
    authors_bib = " and ".join(authors) if authors else "Unknown"
    lines = [f"@article{{{key},"]
    lines.append(f"  title   = {{{title}}},")
    lines.append(f"  author  = {{{authors_bib}}},")
    lines.append(f"  year    = {{{year}}},")
    if source:
        lines.append(f"  journal = {{{source}}},")
    lines.append("}")
    return "\n".join(lines)


def _parse_bibtex(text):
    """Parse BibTeX text into a list of entries."""
    import re
    entries = []
    # Split into entries by @type{...
    entry_pattern = re.compile(r'@(\w+)\s*\{\s*([^,]+),\s*([^@]*?)\n\s*\}', re.DOTALL)
    for match in entry_pattern.finditer(text):
        entry_type = match.group(1).lower()
        key = match.group(2).strip()
        fields_text = match.group(3)
        # Parse fields: key = {value} or key = "value"
        fields = {}
        field_pattern = re.compile(r'(\w+)\s*=\s*[{"]([^}"]+)[}"]\s*,?', re.DOTALL)
        for f_match in field_pattern.finditer(fields_text):
            fields[f_match.group(1).lower()] = f_match.group(2).strip()
        # Convert to citation format
        authors = []
        if "author" in fields:
            authors = [a.strip() for a in fields["author"].split(" and ")]
        entries.append({
            "key": key,
            "type": entry_type,
            "title": fields.get("title", ""),
            "authors": authors,
            "year": fields.get("year", ""),
            "source_info": fields.get("journal", "") or fields.get("booktitle", ""),
        })
    return entries


@app.route("/api/projects/<pid>/bibtex/export", methods=["GET"])
def project_export_bibtex(pid):
    """Export all project citations as a .bib file."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    citations = _load_json(project_path(pid, "citations.json"), default={})
    if not citations:
        return jsonify({"error": "No citations to export."}), 400

    lines = []
    for fname, entry in citations.items():
        # Generate a BibTeX key from author+year
        authors = entry.get("authors", [])
        first_author = authors[0].split()[-1].lower() if authors else "unknown"
        year = entry.get("year", "nd")
        key = f"{first_author}{year}"
        # Sanitize key
        import re
        key = re.sub(r'[^a-zA-Z0-9]', '', key) or "ref"
        lines.append(_citation_to_bibtex(key, entry))

    bib_content = "\n\n".join(lines)
    buf = io.BytesIO(bib_content.encode("utf-8"))
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/x-bibtex",
        as_attachment=True,
        download_name=f"{pid}_bibliography.bib",
    )


@app.route("/api/projects/<pid>/bibtex/import", methods=["POST"])
def project_import_bibtex(pid):
    """Import citations from a BibTeX file (text in request body or file upload)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    # Accept either uploaded file or raw text
    bib_text = ""
    if "file" in request.files:
        bib_text = request.files["file"].read().decode("utf-8", errors="replace")
    else:
        data = request.get_json() or {}
        bib_text = data.get("bibtex", "")

    if not bib_text.strip():
        return jsonify({"error": "No BibTeX content provided."}), 400

    entries = _parse_bibtex(bib_text)
    if not entries:
        return jsonify({"error": "No valid BibTeX entries found."}), 400

    # Store entries as citations — keyed by citation key (not filename, since no PDF)
    cit_path = project_path(pid, "citations.json")
    citations = _load_json(cit_path, default={})
    now = datetime.now().isoformat()
    imported_count = 0
    for entry in entries:
        # Use a synthetic "filename" for BibTeX-only entries: bib:<key>
        synthetic_name = f"bib:{entry['key']}"
        citations[synthetic_name] = {
            "title": entry["title"],
            "authors": entry["authors"],
            "year": entry["year"],
            "source_info": entry["source_info"],
            "generated_at": now,
            "source": "bibtex",
            "bibtex_key": entry["key"],
        }
        imported_count += 1

    _save_json(cit_path, citations)
    return jsonify({"success": True, "imported": imported_count})


# --- DOI Lookup (CrossRef) ---

@app.route("/api/doi-lookup", methods=["GET"])
def doi_lookup():
    """Fetch metadata for a DOI from CrossRef."""
    import requests as req
    doi = request.args.get("doi", "").strip()
    if not doi:
        return jsonify({"error": "DOI parameter required."}), 400

    # Clean DOI (remove URL prefix if present)
    doi = doi.replace("https://doi.org/", "").replace("http://doi.org/", "").strip()

    try:
        resp = req.get(
            f"https://api.crossref.org/works/{doi}",
            headers={"User-Agent": "Scholarium-Research-Assistant (mailto:research@local)"},
            timeout=10,
        )
        if resp.status_code == 404:
            return jsonify({"error": "DOI not found."}), 404
        resp.raise_for_status()
        data = resp.json().get("message", {})

        # Extract authors
        authors = []
        for a in data.get("author", []):
            given = a.get("given", "")
            family = a.get("family", "")
            if family:
                authors.append(f"{given} {family}".strip())

        # Extract year
        year = ""
        issued = data.get("issued", {}).get("date-parts", [[]])
        if issued and issued[0]:
            year = str(issued[0][0])

        # Extract title (first element of title array)
        title = (data.get("title") or [""])[0]

        # Source: journal or book name
        source = (data.get("container-title") or [""])[0]

        return jsonify({
            "doi": doi,
            "title": title,
            "authors": authors,
            "year": year,
            "source_info": source,
            "type": data.get("type", ""),
            "url": data.get("URL", f"https://doi.org/{doi}"),
        })
    except req.Timeout:
        return jsonify({"error": "CrossRef request timed out."}), 504
    except req.ConnectionError:
        return jsonify({"error": "Cannot reach CrossRef. Check your internet connection."}), 503
    except Exception as e:
        return jsonify({"error": f"Lookup failed: {str(e)}"}), 500


@app.route("/api/projects/<pid>/citations/from-doi", methods=["POST"])
def project_create_citation_from_doi(pid):
    """Fetch DOI metadata via CrossRef and save as a citation."""
    import requests as req
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    data = request.get_json() or {}
    doi = data.get("doi", "").strip()
    if not doi:
        return jsonify({"error": "DOI required."}), 400
    doi = doi.replace("https://doi.org/", "").replace("http://doi.org/", "").strip()

    try:
        resp = req.get(
            f"https://api.crossref.org/works/{doi}",
            headers={"User-Agent": "Scholarium-Research-Assistant (mailto:research@local)"},
            timeout=10,
        )
        if resp.status_code == 404:
            return jsonify({"error": "DOI not found."}), 404
        resp.raise_for_status()
        cr = resp.json().get("message", {})

        authors = []
        for a in cr.get("author", []):
            given = a.get("given", "")
            family = a.get("family", "")
            if family:
                authors.append(f"{given} {family}".strip())
        year = ""
        issued = cr.get("issued", {}).get("date-parts", [[]])
        if issued and issued[0]:
            year = str(issued[0][0])

        entry = {
            "title": (cr.get("title") or [""])[0],
            "authors": authors,
            "year": year,
            "source_info": (cr.get("container-title") or [""])[0],
            "doi": doi,
            "source": "crossref",
            "generated_at": datetime.now().isoformat(),
        }

        cit_path = project_path(pid, "citations.json")
        citations = _load_json(cit_path, default={})
        key = f"doi:{doi}"
        citations[key] = entry
        _save_json(cit_path, citations)
        return jsonify({"success": True, "key": key, "entry": entry})
    except Exception as e:
        return jsonify({"error": f"Lookup failed: {str(e)}"}), 500


# --- Notes Search (Full-Text) ---

@app.route("/api/projects/<pid>/notes/search", methods=["GET"])
def project_search_notes(pid):
    """Search notes by title and content."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    q = request.args.get("q", "").strip().lower()
    notes = _load_json(project_path(pid, "notes.json"), default=[])
    if not q:
        return jsonify(notes)
    # Search tag, title, content
    results = []
    for note in notes:
        haystack = (
            (note.get("title", "") + " " +
             note.get("content", "") + " " +
             " ".join(note.get("tags", [])))
            .lower()
        )
        if q in haystack:
            results.append(note)
    return jsonify(results)


# --- Manuscripts ---

def _count_words(html_content):
    """Count words in an HTML content string (strips tags)."""
    if not html_content:
        return 0
    text = re.sub(r'<[^>]+>', ' ', html_content)
    text = html_module.unescape(text)
    tokens = [t for t in text.split() if t.strip()]
    return len(tokens)


def _strip_html(html_content):
    """Strip HTML tags and decode entities, returning plain text."""
    if not html_content:
        return ""
    text = re.sub(r'<[^>]+>', '', html_content)
    return html_module.unescape(text)


def _html_to_paragraphs(html_content):
    """Split HTML content into plain-text paragraphs."""
    if not html_content:
        return []
    # Normalize breaks and closing block tags into newlines
    normalized = re.sub(r'(?i)<br\s*/?>', '\n', html_content)
    normalized = re.sub(r'(?i)</p\s*>', '\n', normalized)
    normalized = re.sub(r'(?i)</div\s*>', '\n', normalized)
    normalized = re.sub(r'(?i)</li\s*>', '\n', normalized)
    normalized = re.sub(r'(?i)</h[1-6]\s*>', '\n', normalized)
    # Strip remaining tags
    text = re.sub(r'<[^>]+>', '', normalized)
    text = html_module.unescape(text)
    paragraphs = [p.strip() for p in text.split("\n")]
    return [p for p in paragraphs if p]


def _latin1(text):
    """Encode text to latin-1 with replacement for fpdf2 compatibility."""
    return (text or "").encode("latin-1", "replace").decode("latin-1")


def _manuscript_summary(m):
    """Return a manuscript dict with a total_words summary field."""
    total = sum(_count_words(c.get("content", "")) for c in m.get("chapters", []))
    return {
        "id": m.get("id"),
        "title": m.get("title"),
        "citation_style": m.get("citation_style", "apa"),
        "chapter_count": len(m.get("chapters", [])),
        "total_words": total,
        "citations_used": m.get("citations_used", []),
        "created_at": m.get("created_at"),
        "updated_at": m.get("updated_at"),
    }


def _log_writing(pid, mid, delta_words):
    """Record positive word deltas to writing_log.json under today's date."""
    log_path = project_path(pid, "writing_log.json")
    log = _load_json(log_path, default={})
    today = datetime.now().strftime("%Y-%m-%d")
    entry = log.get(today) or {"words_added": 0, "manuscripts_touched": []}
    if delta_words and delta_words > 0:
        entry["words_added"] = int(entry.get("words_added", 0)) + int(delta_words)
    touched = entry.get("manuscripts_touched") or []
    if mid and mid not in touched:
        touched.append(mid)
    entry["manuscripts_touched"] = touched
    log[today] = entry
    _save_json(log_path, log)


def _snapshot_version(pid, mid, cid, old_content):
    """Append a snapshot of old chapter content to versions.json, keeping last 20."""
    ver_path = project_path(pid, "versions.json")
    versions = _load_json(ver_path, default={})
    key = f"{mid}:{cid}"
    entries = versions.get(key, [])
    entries.append({
        "timestamp": datetime.now().isoformat(),
        "content": old_content or "",
        "words": _count_words(old_content or ""),
    })
    if len(entries) > 20:
        entries = entries[-20:]
    versions[key] = entries
    _save_json(ver_path, versions)


@app.route("/api/projects/<pid>/manuscripts", methods=["GET"])
def project_list_manuscripts(pid):
    """List all manuscripts with summary info."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    manuscripts = _load_json(project_path(pid, "manuscripts.json"), default=[])
    return jsonify([_manuscript_summary(m) for m in manuscripts])


@app.route("/api/projects/<pid>/manuscripts", methods=["POST"])
def project_create_manuscript(pid):
    """Create a new manuscript."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "Title is required."}), 400

    now = datetime.now().isoformat()
    manuscript = {
        "id": uuid.uuid4().hex[:12],
        "title": title,
        "chapters": [],
        "citations_used": [],
        "citation_style": data.get("citation_style", "apa"),
        "created_at": now,
        "updated_at": now,
    }

    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    manuscripts.append(manuscript)
    _save_json(path, manuscripts)
    return jsonify(manuscript), 201


@app.route("/api/projects/<pid>/manuscripts/<mid>", methods=["GET"])
def project_get_manuscript(pid, mid):
    """Return a manuscript with all chapter content."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    manuscripts = _load_json(project_path(pid, "manuscripts.json"), default=[])
    for m in manuscripts:
        if m.get("id") == mid:
            result = dict(m)
            result["total_words"] = sum(_count_words(c.get("content", "")) for c in m.get("chapters", []))
            return jsonify(result)
    return jsonify({"error": "Manuscript not found"}), 404


@app.route("/api/projects/<pid>/manuscripts/<mid>", methods=["PUT"])
def project_update_manuscript(pid, mid):
    """Update manuscript title or citation_style."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    for m in manuscripts:
        if m.get("id") == mid:
            if "title" in data and data["title"] is not None:
                m["title"] = data["title"]
            if "citation_style" in data and data["citation_style"] is not None:
                m["citation_style"] = data["citation_style"]
            if "citations_used" in data and isinstance(data["citations_used"], list):
                m["citations_used"] = data["citations_used"]
            m["updated_at"] = datetime.now().isoformat()
            _save_json(path, manuscripts)
            return jsonify(m)
    return jsonify({"error": "Manuscript not found"}), 404


@app.route("/api/projects/<pid>/manuscripts/<mid>", methods=["DELETE"])
def project_delete_manuscript(pid, mid):
    """Delete a manuscript by id."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    filtered = [m for m in manuscripts if m.get("id") != mid]
    if len(filtered) == len(manuscripts):
        return jsonify({"error": "Manuscript not found"}), 404
    _save_json(path, filtered)
    return jsonify({"success": True})


@app.route("/api/projects/<pid>/manuscripts/<mid>/chapters", methods=["POST"])
def project_add_chapter(pid, mid):
    """Add a chapter to a manuscript."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "Chapter title is required."}), 400

    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    for m in manuscripts:
        if m.get("id") == mid:
            now = datetime.now().isoformat()
            chapters = m.setdefault("chapters", [])
            next_order = max((c.get("order", 0) for c in chapters), default=-1) + 1
            chapter = {
                "id": uuid.uuid4().hex[:12],
                "title": title,
                "content": data.get("content", ""),
                "order": next_order,
                "created_at": now,
                "updated_at": now,
            }
            chapters.append(chapter)
            m["updated_at"] = now
            _save_json(path, manuscripts)
            return jsonify(chapter), 201
    return jsonify({"error": "Manuscript not found"}), 404


@app.route("/api/projects/<pid>/manuscripts/<mid>/chapters/<cid>", methods=["PUT"])
def project_update_chapter(pid, mid, cid):
    """Update a chapter: title, content, order. Logs word diff and snapshots."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    for m in manuscripts:
        if m.get("id") != mid:
            continue
        for c in m.get("chapters", []):
            if c.get("id") != cid:
                continue

            old_content = c.get("content", "")
            content_changed = "content" in data and data["content"] is not None and data["content"] != old_content

            if content_changed:
                _snapshot_version(pid, mid, cid, old_content)
                old_words = _count_words(old_content)
                new_words = _count_words(data["content"])
                c["content"] = data["content"]
                delta = new_words - old_words
                if delta > 0:
                    _log_writing(pid, mid, delta)
                else:
                    # Still mark manuscript as touched today
                    _log_writing(pid, mid, 0)

            if "title" in data and data["title"] is not None:
                c["title"] = data["title"]
            if "order" in data and data["order"] is not None:
                try:
                    c["order"] = int(data["order"])
                except (TypeError, ValueError):
                    pass

            c["updated_at"] = datetime.now().isoformat()
            m["updated_at"] = c["updated_at"]
            _save_json(path, manuscripts)
            return jsonify(c)
        return jsonify({"error": "Chapter not found"}), 404
    return jsonify({"error": "Manuscript not found"}), 404


@app.route("/api/projects/<pid>/manuscripts/<mid>/chapters/<cid>", methods=["DELETE"])
def project_delete_chapter(pid, mid, cid):
    """Remove a chapter from a manuscript."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    for m in manuscripts:
        if m.get("id") != mid:
            continue
        chapters = m.get("chapters", [])
        new_chapters = [c for c in chapters if c.get("id") != cid]
        if len(new_chapters) == len(chapters):
            return jsonify({"error": "Chapter not found"}), 404
        m["chapters"] = new_chapters
        m["updated_at"] = datetime.now().isoformat()
        _save_json(path, manuscripts)
        return jsonify({"success": True})
    return jsonify({"error": "Manuscript not found"}), 404


@app.route("/api/projects/<pid>/manuscripts/<mid>/versions/<cid>", methods=["GET"])
def project_list_chapter_versions(pid, mid, cid):
    """List saved versions for a chapter."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    versions = _load_json(project_path(pid, "versions.json"), default={})
    key = f"{mid}:{cid}"
    return jsonify(versions.get(key, []))


@app.route("/api/projects/<pid>/manuscripts/<mid>/chapters/<cid>/restore", methods=["POST"])
def project_restore_chapter_version(pid, mid, cid):
    """Restore chapter content from a prior version (snapshots current first)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    if "version_index" not in data:
        return jsonify({"error": "version_index is required."}), 400
    try:
        idx = int(data["version_index"])
    except (TypeError, ValueError):
        return jsonify({"error": "version_index must be an integer."}), 400

    ver_path = project_path(pid, "versions.json")
    versions = _load_json(ver_path, default={})
    key = f"{mid}:{cid}"
    entries = versions.get(key, [])
    if idx < 0 or idx >= len(entries):
        return jsonify({"error": "version_index out of range."}), 400

    target = entries[idx]
    restored_content = target.get("content", "")

    path = project_path(pid, "manuscripts.json")
    manuscripts = _load_json(path, default=[])
    for m in manuscripts:
        if m.get("id") != mid:
            continue
        for c in m.get("chapters", []):
            if c.get("id") != cid:
                continue
            old_content = c.get("content", "")
            _snapshot_version(pid, mid, cid, old_content)
            old_words = _count_words(old_content)
            new_words = _count_words(restored_content)
            c["content"] = restored_content
            c["updated_at"] = datetime.now().isoformat()
            m["updated_at"] = c["updated_at"]
            _save_json(path, manuscripts)
            delta = new_words - old_words
            if delta > 0:
                _log_writing(pid, mid, delta)
            else:
                _log_writing(pid, mid, 0)
            return jsonify(c)
        return jsonify({"error": "Chapter not found"}), 404
    return jsonify({"error": "Manuscript not found"}), 404


def _export_manuscript_pdf(manuscript, bibliography):
    """Build a PDF for a manuscript and return a Flask send_file response."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title page
    pdf.add_page()
    pdf.ln(60)
    pdf.set_font("Helvetica", "B", 22)
    pdf.multi_cell(0, 12, _latin1(manuscript.get("title", "Untitled")), align="C")
    pdf.ln(6)
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, _latin1(datetime.now().strftime("%B %d, %Y")),
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_text_color(0, 0, 0)

    # Chapters, one per page
    chapters = sorted(manuscript.get("chapters", []), key=lambda c: c.get("order", 0))
    for ch in chapters:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 10, _latin1(ch.get("title", "Untitled Chapter")))
        pdf.ln(4)
        pdf.set_font("Helvetica", "", 12)
        body = _strip_html(ch.get("content", ""))
        body = _latin1(body)
        if body.strip():
            pdf.multi_cell(0, 7, body)

    # Bibliography page
    if bibliography:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Bibliography", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        pdf.set_font("Helvetica", "", 11)
        for line in bibliography:
            pdf.multi_cell(0, 6, _latin1(line))
            pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    safe_title = re.sub(r'[^A-Za-z0-9_-]+', '_', manuscript.get("title", "manuscript"))[:60] or "manuscript"
    return send_file(buf, mimetype="application/pdf", as_attachment=True,
                     download_name=f"{safe_title}.pdf")


def _export_manuscript_docx(manuscript, bibliography):
    """Build a DOCX for a manuscript and return a Flask send_file response."""
    from docx import Document

    doc = Document()

    # Title page
    doc.add_heading(manuscript.get("title", "Untitled"), level=0)
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_page_break()

    chapters = sorted(manuscript.get("chapters", []), key=lambda c: c.get("order", 0))
    for idx, ch in enumerate(chapters):
        doc.add_heading(ch.get("title", "Untitled Chapter"), level=1)
        for para in _html_to_paragraphs(ch.get("content", "")):
            doc.add_paragraph(para)
        if idx < len(chapters) - 1:
            doc.add_page_break()

    if bibliography:
        doc.add_page_break()
        doc.add_heading("Bibliography", level=1)
        for line in bibliography:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_title = re.sub(r'[^A-Za-z0-9_-]+', '_', manuscript.get("title", "manuscript"))[:60] or "manuscript"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{safe_title}.docx",
    )


@app.route("/api/projects/<pid>/manuscripts/<mid>/export", methods=["GET"])
def project_export_manuscript(pid, mid):
    """Export a manuscript as PDF or DOCX, including a bibliography."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    fmt = (request.args.get("format", "pdf") or "pdf").lower()
    style = (request.args.get("style") or "").lower()

    manuscripts = _load_json(project_path(pid, "manuscripts.json"), default=[])
    manuscript = next((m for m in manuscripts if m.get("id") == mid), None)
    if not manuscript:
        return jsonify({"error": "Manuscript not found"}), 404

    if not style:
        style = (manuscript.get("citation_style") or "apa").lower()

    # Build bibliography from citations_used
    citations = _load_json(project_path(pid, "citations.json"), default={})
    bibliography = []
    for key in manuscript.get("citations_used", []) or []:
        entry = citations.get(key)
        if not entry:
            continue
        formatted = _format_citation(entry, style)
        if formatted:
            bibliography.append(formatted)
    bibliography.sort(key=lambda s: s.lower())

    if fmt == "pdf":
        return _export_manuscript_pdf(manuscript, bibliography)
    if fmt == "docx":
        return _export_manuscript_docx(manuscript, bibliography)
    return jsonify({"error": "Unsupported format. Use pdf or docx."}), 400


@app.route("/api/projects/<pid>/writing-streak", methods=["GET"])
def project_writing_streak(pid):
    """Compute current writing streak (consecutive days with >100 words)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    log = _load_json(project_path(pid, "writing_log.json"), default={})
    today = datetime.now().date()
    today_key = today.strftime("%Y-%m-%d")
    today_entry = log.get(today_key) or {}
    today_words = int(today_entry.get("words_added", 0) or 0)

    streak = 0
    days = []
    # Start from today; if today has >100, count it. Otherwise start from yesterday.
    if today_words > 100:
        cursor = today
    else:
        cursor = today - timedelta(days=1)

    while True:
        key = cursor.strftime("%Y-%m-%d")
        entry = log.get(key) or {}
        words = int(entry.get("words_added", 0) or 0)
        if words > 100:
            streak += 1
            days.append({"date": key, "words_added": words})
            cursor = cursor - timedelta(days=1)
        else:
            break

    return jsonify({
        "streak": streak,
        "today_words": today_words,
        "days": days,
    })


# --- Collections ---

@app.route("/api/projects/<pid>/collections", methods=["GET"])
def project_get_collections(pid):
    """List all collections for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    col_path = project_path(pid, "collections.json")
    return jsonify(_load_json(col_path, default=[]))


@app.route("/api/projects/<pid>/collections", methods=["POST"])
def project_create_collection(pid):
    """Create a new collection in a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    if not data.get("name"):
        return jsonify({"error": "Collection name is required."}), 400

    now = datetime.now().isoformat()
    collection = {
        "id": uuid.uuid4().hex[:12],
        "name": data["name"],
        "documents": data.get("documents", []),
        "created_at": now,
        "updated_at": now,
    }

    col_path = project_path(pid, "collections.json")
    collections = _load_json(col_path, default=[])
    collections.append(collection)
    _save_json(col_path, collections)
    return jsonify(collection), 201


@app.route("/api/projects/<pid>/collections/<cid>", methods=["PUT"])
def project_update_collection(pid, cid):
    """Update a collection."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    col_path = project_path(pid, "collections.json")
    collections = _load_json(col_path, default=[])
    for col in collections:
        if col["id"] == cid:
            if "name" in data:
                col["name"] = data["name"]
            if "documents" in data:
                col["documents"] = data["documents"]
            col["updated_at"] = datetime.now().isoformat()
            _save_json(col_path, collections)
            return jsonify(col)
    return jsonify({"error": "Collection not found"}), 404


@app.route("/api/projects/<pid>/collections/<cid>", methods=["DELETE"])
def project_delete_collection(pid, cid):
    """Delete a collection."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    col_path = project_path(pid, "collections.json")
    collections = _load_json(col_path, default=[])
    filtered = [c for c in collections if c["id"] != cid]
    if len(filtered) == len(collections):
        return jsonify({"error": "Collection not found"}), 404
    _save_json(col_path, filtered)
    return jsonify({"success": True})


# --- Journal ---

def _compute_today_stats(pid):
    """Compute auto_stats for today by counting Q&A log entries and reading status."""
    today_str = datetime.now().strftime("%Y-%m-%d")
    qa_count = 0
    qa_path = project_path(pid, "qa_log.json")
    qa_log = _load_json(qa_path, default=[])
    for entry in qa_log:
        ts = entry.get("timestamp", "")
        if ts.startswith(today_str):
            qa_count += 1

    docs_viewed = []
    rs_path = project_path(pid, "reading_status.json")
    statuses = _load_json(rs_path, default={})
    for fname, info in statuses.items():
        last_opened = info.get("last_opened", "")
        if last_opened and last_opened.startswith(today_str):
            docs_viewed.append(fname)

    return {"qa_count": qa_count, "docs_viewed": docs_viewed}


@app.route("/api/projects/<pid>/journal", methods=["GET"])
def project_get_journal(pid):
    """Return all journal entries for a project, newest first."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    j_path = project_path(pid, "journal.json")
    entries = _load_json(j_path, default=[])
    entries.sort(key=lambda e: e.get("created_at", ""), reverse=True)
    return jsonify(entries)


@app.route("/api/projects/<pid>/journal/today", methods=["GET"])
def project_get_journal_today(pid):
    """Get or create today's journal entry with auto-populated stats."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    today_str = datetime.now().strftime("%Y-%m-%d")
    j_path = project_path(pid, "journal.json")
    entries = _load_json(j_path, default=[])

    for entry in entries:
        if entry.get("date") == today_str:
            entry["auto_stats"] = _compute_today_stats(pid)
            _save_json(j_path, entries)
            return jsonify(entry)

    now = datetime.now().isoformat()
    entry = {
        "id": uuid.uuid4().hex[:12],
        "date": today_str,
        "content": "",
        "auto_stats": _compute_today_stats(pid),
        "created_at": now,
        "updated_at": now,
    }
    entries.append(entry)
    _save_json(j_path, entries)
    return jsonify(entry)


@app.route("/api/projects/<pid>/journal/<jid>", methods=["PUT"])
def project_update_journal(pid, jid):
    """Update a journal entry's content."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    j_path = project_path(pid, "journal.json")
    entries = _load_json(j_path, default=[])
    for entry in entries:
        if entry["id"] == jid:
            if "content" in data:
                entry["content"] = data["content"]
            entry["updated_at"] = datetime.now().isoformat()
            _save_json(j_path, entries)
            return jsonify(entry)
    return jsonify({"error": "Journal entry not found"}), 404


# --- Literature Matrix (background-worker based) ---

# In-memory job tracker for matrix generation jobs (keyed by project id)
# Each entry: {"running": bool, "total": N, "done": M, "errors": [], "started_at": ISO, "finished_at": ISO|None}
_matrix_jobs: dict[str, dict] = {}
_matrix_jobs_lock = threading.Lock()

# Generic job registry for all long-running background tasks.
# Keyed by project id, then by job id.
#   type ∈ {"summarize", "citation", "matrix_entry"}
#   status ∈ {"running", "done", "error"}
_jobs: dict[str, dict[str, dict]] = {}
_jobs_lock = threading.Lock()


def _add_job(pid: str, job_id: str, job_type: str, target: str):
    with _jobs_lock:
        _jobs.setdefault(pid, {})[job_id] = {
            "id": job_id,
            "type": job_type,
            "target": target,
            "status": "running",
            "started_at": datetime.now().isoformat(),
            "finished_at": None,
            "error": None,
        }


def _finish_job(pid: str, job_id: str, error: str | None = None):
    with _jobs_lock:
        job = _jobs.get(pid, {}).get(job_id)
        if job:
            job["status"] = "error" if error else "done"
            job["finished_at"] = datetime.now().isoformat()
            if error:
                job["error"] = str(error)


def _list_jobs(pid: str, running_only: bool = False) -> list[dict]:
    """Return jobs for a project. Removes finished jobs older than 60s."""
    cutoff = (datetime.now() - timedelta(seconds=60)).isoformat()
    with _jobs_lock:
        project_jobs = _jobs.get(pid, {})
        # Prune old finished jobs
        stale_ids = [
            jid for jid, j in project_jobs.items()
            if j["status"] != "running" and (j.get("finished_at") or "") < cutoff
        ]
        for jid in stale_ids:
            project_jobs.pop(jid, None)
        out = list(project_jobs.values())
    if running_only:
        out = [j for j in out if j["status"] == "running"]
    return out


def _job_exists_running(pid: str, job_id: str) -> bool:
    with _jobs_lock:
        job = _jobs.get(pid, {}).get(job_id)
        return bool(job and job["status"] == "running")


def _matrix_worker(pid: str, docs_to_process: list):
    """Background thread — processes each doc and appends to matrix file as it goes."""
    import fitz
    import requests as req
    from respondent import MODEL, OLLAMA_URL

    m_path = project_path(pid, "literature_matrix.json")
    docs_dir = project_docs_dir(pid)

    for fname in docs_to_process:
        fpath = os.path.join(docs_dir, fname)

        entry = None
        error = None

        try:
            doc = fitz.open(fpath)
            text = ""
            for i, page in enumerate(doc):
                if i >= 4:
                    break
                text += page.get_text()
            doc.close()
            text = text[:8000]

            prompt = (
                "Extract the following metadata from this academic paper text. "
                "Respond in JSON format with these exact keys: "
                "title (string), year (string), methodology (string, brief description of research method), "
                "findings (string, key findings summary), sample_size (string, e.g. 'N=150' or 'N/A' if not applicable).\n\n"
                f"--- DOCUMENT TEXT ---\n{text}"
            )

            response = req.post(
                OLLAMA_URL,
                json={
                    "model": MODEL,
                    "prompt": prompt,
                    "system": "You are an academic paper metadata extractor. Return valid JSON only.",
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 512},
                },
                timeout=180,
            )
            response.raise_for_status()
            raw = response.json().get("response", "").strip()

            json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
            if json_match:
                metadata = json.loads(json_match.group())
            else:
                metadata = json.loads(raw)

            entry = {
                "filename": fname,
                "title": metadata.get("title", ""),
                "year": str(metadata.get("year", "")),
                "methodology": metadata.get("methodology", ""),
                "findings": metadata.get("findings", ""),
                "sample_size": metadata.get("sample_size", ""),
            }
        except Exception as e:
            error = {"filename": fname, "error": str(e)}

        # Persist this iteration's result immediately
        matrix = _load_json(m_path, default={"generated_at": None, "entries": []})
        if entry:
            # Skip duplicates in case of overlapping runs
            if not any(e.get("filename") == fname for e in matrix["entries"]):
                matrix["entries"].append(entry)
        matrix["generated_at"] = datetime.now().isoformat()
        _save_json(m_path, matrix)

        with _matrix_jobs_lock:
            job = _matrix_jobs.get(pid)
            if job:
                job["done"] += 1
                if error:
                    job["errors"].append(error)
                job["current"] = fname

    # Mark job finished
    with _matrix_jobs_lock:
        job = _matrix_jobs.get(pid)
        if job:
            job["running"] = False
            job["finished_at"] = datetime.now().isoformat()
            job["current"] = None


@app.route("/api/projects/<pid>/matrix", methods=["GET"])
def project_get_matrix(pid):
    """Return the cached literature matrix for a project, with relevance scores merged in."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    m_path = project_path(pid, "literature_matrix.json")
    matrix = _load_json(m_path, default={"generated_at": None, "entries": []})

    # Merge relevance scores into each entry
    relevance_data = _load_json(project_path(pid, "relevance.json"), default={"scores": {}})
    scores_map = relevance_data.get("scores", {})
    for entry in matrix.get("entries", []):
        rel = scores_map.get(entry.get("filename"))
        if rel:
            entry["relevance_score"] = rel.get("score")

    return jsonify(matrix)


@app.route("/api/projects/<pid>/jobs", methods=["GET"])
def project_list_jobs(pid):
    """Return all background jobs for a project (summarize / matrix_entry / etc.).
    By default shows running + recently-finished jobs (last 60s)."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404
    running_only = request.args.get("running", "").lower() == "true"
    return jsonify(_list_jobs(pid, running_only=running_only))


@app.route("/api/projects/<pid>/matrix/status", methods=["GET"])
def project_matrix_status(pid):
    """Return the current matrix generation job status for a project."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    with _matrix_jobs_lock:
        job = _matrix_jobs.get(pid)
        if not job:
            return jsonify({"running": False, "total": 0, "done": 0, "errors": [], "current": None})
        # Return a copy
        return jsonify({
            "running": job.get("running", False),
            "total": job.get("total", 0),
            "done": job.get("done", 0),
            "errors": list(job.get("errors", [])),
            "current": job.get("current"),
            "started_at": job.get("started_at"),
            "finished_at": job.get("finished_at"),
        })


@app.route("/api/projects/<pid>/matrix/generate", methods=["POST"])
def project_generate_matrix(pid):
    """Kick off background matrix generation for docs not already in the matrix.
    Returns immediately; poll /matrix/status and /matrix for progress + results."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    # Don't start a second job for this project
    with _matrix_jobs_lock:
        current_job = _matrix_jobs.get(pid)
        if current_job and current_job.get("running"):
            return jsonify({
                "message": "Generation already running.",
                "status": {
                    "running": True,
                    "total": current_job.get("total", 0),
                    "done": current_job.get("done", 0),
                    "current": current_job.get("current"),
                },
            }), 202

    m_path = project_path(pid, "literature_matrix.json")
    matrix = _load_json(m_path, default={"generated_at": None, "entries": []})
    existing_filenames = {e["filename"] for e in matrix["entries"]}

    docs_dir = project_docs_dir(pid)
    docs_to_process = []
    if os.path.isdir(docs_dir):
        for fname in sorted(os.listdir(docs_dir)):
            if not fname.lower().endswith(".pdf"):
                continue
            if fname not in existing_filenames:
                docs_to_process.append(fname)

    if not docs_to_process:
        return jsonify({"message": "All documents already in matrix.", "status": {"running": False, "total": 0, "done": 0}})

    # Initialize job state
    with _matrix_jobs_lock:
        _matrix_jobs[pid] = {
            "running": True,
            "total": len(docs_to_process),
            "done": 0,
            "errors": [],
            "current": docs_to_process[0] if docs_to_process else None,
            "started_at": datetime.now().isoformat(),
            "finished_at": None,
        }

    # Spawn background thread (daemon=True so it dies with the server)
    t = threading.Thread(
        target=_matrix_worker,
        args=(pid, docs_to_process),
        daemon=True,
    )
    t.start()

    return jsonify({
        "message": f"Generation started for {len(docs_to_process)} document(s). Poll /matrix/status for progress.",
        "status": {
            "running": True,
            "total": len(docs_to_process),
            "done": 0,
            "current": docs_to_process[0],
        },
    }), 202


# --- Answer Export ---

@app.route("/api/projects/<pid>/answer/export", methods=["POST"])
def project_export_answer(pid):
    """Export a Q&A answer as PDF or DOCX."""
    if not _get_project(pid):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    fmt = data.get("format", "pdf").lower()
    question = data.get("question", "")
    answer = data.get("answer", "")
    scores = data.get("scores", {})
    model = data.get("model", "")
    timestamp = data.get("timestamp", "")

    if not answer:
        return jsonify({"error": "No answer to export"}), 400

    if fmt == "pdf":
        return _export_answer_pdf(question, answer, scores, model, timestamp)
    elif fmt == "docx":
        return _export_answer_docx(question, answer, scores, model, timestamp)
    else:
        return jsonify({"error": "Unsupported format. Use 'pdf' or 'docx'."}), 400


# ===========================================================================
# Export helper functions
# ===========================================================================

def _export_pdf(note: dict):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, note["title"], new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Date
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(120, 120, 120)
    created = note.get("created_at", "")[:19].replace("T", " ")
    pdf.cell(0, 6, f"Created: {created}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Content
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 12)
    content = note["content"].encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 7, content)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    filename = f"{note['title'][:50].replace(' ', '_')}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True,
                     download_name=filename)


def _export_docx(note: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    doc.add_heading(note["title"], level=1)

    # Date
    date_para = doc.add_paragraph()
    created = note.get("created_at", "")[:19].replace("T", " ")
    run = date_para.add_run(f"Created: {created}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Content
    for paragraph_text in note["content"].split("\n"):
        doc.add_paragraph(paragraph_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{note['title'][:50].replace(' ', '_')}.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


def _export_answer_pdf(question, answer, scores, model, timestamp):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Scholarium - Research Assistant - Q&A Export", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Metadata
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(120, 120, 120)
    if timestamp:
        pdf.cell(0, 5, f"Date: {timestamp[:19].replace('T', ' ')}", new_x="LMARGIN", new_y="NEXT")
    if model:
        pdf.cell(0, 5, f"Model: {model}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Question
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Question:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    q_text = question.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 7, q_text)
    pdf.ln(4)

    # Answer
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Answer:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    a_text = answer.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 6, a_text)
    pdf.ln(4)

    # Scores
    if scores:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Evaluation Scores:", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for dim, sdata in scores.items():
            if isinstance(sdata, dict):
                score_val = sdata.get("score", "N/A")
                explanation = sdata.get("explanation", "")
                label = dim.replace("_", " ").title()
                pct = f"{int(float(score_val) * 100)}%" if isinstance(score_val, (int, float)) else str(score_val)
                line = f"{label}: {pct} - {explanation}"
                line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 6, line)
                pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/pdf", as_attachment=True,
                     download_name="research_answer.pdf")


def _export_answer_docx(question, answer, scores, model, timestamp):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Scholarium - Research Assistant - Q&A Export", level=1)

    # Metadata
    if timestamp or model:
        meta = doc.add_paragraph()
        if timestamp:
            run = meta.add_run(f"Date: {timestamp[:19].replace('T', ' ')}   ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
        if model:
            run = meta.add_run(f"Model: {model}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)

    # Question
    doc.add_heading("Question", level=2)
    doc.add_paragraph(question)

    # Answer
    doc.add_heading("Answer", level=2)
    for para in answer.split("\n"):
        doc.add_paragraph(para)

    # Scores
    if scores:
        doc.add_heading("Evaluation Scores", level=2)
        for dim, sdata in scores.items():
            if isinstance(sdata, dict):
                score_val = sdata.get("score", "N/A")
                explanation = sdata.get("explanation", "")
                label = dim.replace("_", " ").title()
                pct = f"{int(float(score_val) * 100)}%" if isinstance(score_val, (int, float)) else str(score_val)
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: {pct}")
                run.bold = True
                p.add_run(f" - {explanation}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="research_answer.docx",
    )


# ===========================================================================
# Startup
# ===========================================================================

def open_browser():
    """Open the browser after a short delay."""
    import time
    time.sleep(1.5)
    webbrowser.open("http://localhost:8080")


if __name__ == "__main__":
    # Ensure projects directory exists
    os.makedirs(PROJECTS_DIR, exist_ok=True)

    # Migrate old single-project data if present
    print("Checking for legacy data to migrate...")
    _migrate_old_data()

    # Index documents for every existing project
    projects = _load_projects()
    for proj in projects:
        pid = proj["id"]
        print(f"Indexing project '{proj['name']}' ({pid})...")
        ev = get_evaluator(pid)
        idx_result = ev.index_documents()
        print(f"  Files processed: {idx_result['files_processed']}, "
              f"Chunks added: {idx_result['total_chunks']}, "
              f"Skipped: {idx_result['files_skipped']}")
        status = ev.get_status()
        print(f"  Total documents: {status['documents_indexed']}, "
              f"Total chunks: {status['total_chunks']}")

    if not projects:
        print("\n  No projects found. Create one via the web UI.\n")

    # Open browser in background thread (skip when running in Docker / headless)
    if not os.environ.get("SCHOLARIUM_NO_BROWSER"):
        threading.Thread(target=open_browser, daemon=True).start()

    print("Starting server at http://localhost:8080")
    app.run(host="0.0.0.0", port=8080, debug=False)
